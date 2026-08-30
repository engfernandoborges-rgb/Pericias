import os
import openpyxl
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, hex_color):
    """Define a cor de fundo de uma célula (Shading)."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Define as margens internas (padding) de uma célula em dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders_to_none(table):
    """Remove completamente as bordas de uma tabela."""
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def add_images_side_by_side(doc, img1_path, img2_path, caption1, caption2):
    """Insere duas imagens lado a lado em uma tabela sem bordas com legendas."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Remove as bordas para visualização limpa
    set_table_borders_to_none(table)
    
    # Define larguras fixas de 3 polegadas para cada lado
    for row in table.rows:
        row.cells[0].width = Inches(3.0)
        row.cells[1].width = Inches(3.0)
        
    # Célula 1 - Visão Geral
    cell1 = table.cell(0, 0)
    p1 = cell1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(img1_path):
        try:
            run1 = p1.add_run()
            run1.add_picture(img1_path, width=Inches(2.8))
            p1_cap = cell1.add_paragraph()
            p1_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run1_cap = p1_cap.add_run(caption1)
            run1_cap.font.size = Pt(9)
            run1_cap.font.italic = True
            run1_cap.font.name = 'Calibri'
        except Exception as e:
            p1.add_run(f"🚨 Erro ao carregar Foto 1: {e}").font.size = Pt(9)
    else:
        p1.add_run(f"📷 Foto 1 Ausente: {os.path.basename(img1_path)}").font.size = Pt(9)
        
    # Célula 2 - Detalhe
    cell2 = table.cell(0, 1)
    p2 = cell2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(img2_path):
        try:
            run2 = p2.add_run()
            run2.add_picture(img2_path, width=Inches(2.8))
            p2_cap = cell2.add_paragraph()
            p2_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2_cap = p2_cap.add_run(caption2)
            run2_cap.font.size = Pt(9)
            run2_cap.font.italic = True
            run2_cap.font.name = 'Calibri'
        except Exception as e:
            p2.add_run(f"🚨 Erro ao carregar Foto 2: {e}").font.size = Pt(9)
    else:
        p2.add_run(f"📷 Foto 2 Ausente: {os.path.basename(img2_path)}").font.size = Pt(9)

def format_run(run, font_name='Calibri', size_pt=11, bold=False, italic=False, color_rgb=(0,0,0)):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(*color_rgb)

def generate_docx_laudo(excel_path, fotos_dir, output_docx_path):
    # Carrega a planilha Excel
    if not os.path.exists(excel_path):
        raise FileNotFoundError(f"Planilha {excel_path} não encontrada.")
        
    wb = openpyxl.load_workbook(excel_path, data_only=True)
    
    # --- 1. LER DADOS DE IDENTIFICAÇÃO ---
    ws_ident = wb['1. Identificação e Vistoria']
    processo = ws_ident['B5'].value or "NÃO ESPECIFICADO"
    comarca = ws_ident['E5'].value or "NÃO ESPECIFICADO"
    requerente = ws_ident['B6'].value or "NÃO ESPECIFICADO"
    requerido = ws_ident['E6'].value or "NÃO ESPECIFICADO"
    tipo_pericia = ws_ident['B7'].value or "Vistoria Cautelar"
    data_vistoria = ws_ident['B10'].value or "___/___/______"
    hora_vistoria = ws_ident['E10'].value or "__:__"
    endereco = ws_ident['B11'].value or "NÃO ESPECIFICADO"
    latitude = ws_ident['B12'].value or 0.0
    longitude = ws_ident['E12'].value or 0.0
    
    # Formata data se for objeto datetime
    if hasattr(data_vistoria, 'strftime'):
        data_vistoria = data_vistoria.strftime("%d/%m/%Y")
        
    # --- 2. LER REGISTROS DE DANOS ---
    ws_danos = wb['2. Registro de Danos']
    danos = []
    r_idx = 8
    while True:
        ambiente = ws_danos.cell(r_idx, 1).value
        elemento = ws_danos.cell(r_idx, 2).value
        dano_tipo = ws_danos.cell(r_idx, 3).value
        
        # Streak check de fim de arquivo
        if ambiente is None and elemento is None and dano_tipo is None:
            empty_streak = True
            for check_r in range(r_idx, r_idx + 5):
                if any(ws_danos.cell(check_r, c).value is not None for c in range(1, 11)):
                    empty_streak = False
                    r_idx = check_r
                    break
            if empty_streak:
                break
                
        # Captura registro
        foto_nomes = ws_danos.cell(r_idx, 7).value or ""
        lista_fotos = [f.strip() for f in str(foto_nomes).split(';') if f.strip()]
        
        danos.append({
            'linha': r_idx,
            'ambiente': ambiente or "Não informado",
            'elemento': elemento or "Não informado",
            'dano': dano_tipo or "Não informado",
            'descricao': ws_danos.cell(r_idx, 4).value or "Sem descrição técnica.",
            'medicoes': ws_danos.cell(r_idx, 5).value or "Nenhuma medição registrada.",
            'observacoes': ws_danos.cell(r_idx, 6).value or "Sem observações complementares.",
            'local_exato': ws_danos.cell(r_idx, 8).value or "Área não especificada",
            'fotos': lista_fotos
        })
        r_idx += 1
        
    # --- 3. CRIAR DOCUMENTO WORD ---
    doc = Document()
    
    # Margens Padrão de 2,5 cm (Aprox 1 polegada)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # --- CAPA ---
    # Cabeçalho da capa
    p_topo = doc.add_paragraph()
    p_topo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_topo = p_topo.add_run("PODER JUDICIÁRIO DO ESTADO\nJUÍZO DE DIREITO DA COMARCA DE " + str(comarca).upper())
    format_run(run_topo, font_name='Calibri', size_pt=10, bold=True, color_rgb=(100, 100, 100))
    
    for _ in range(5):
        doc.add_paragraph()
        
    # Título Principal
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tit = p_titulo.add_run("LAUDO PERICIAL DE ENGENHARIA CIVIL")
    format_run(run_tit, font_name='Calibri', size_pt=18, bold=True, color_rgb=(47, 84, 150)) # Azul #2F5496
    
    # Subtítulo
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run(f"PROCESSO Nº: {processo}\nAÇÃO: {tipo_pericia}")
    format_run(run_sub, font_name='Calibri', size_pt=12, bold=True, italic=True, color_rgb=(85, 85, 85))
    
    for _ in range(7):
        doc.add_paragraph()
        
    # Partes
    p_partes = doc.add_paragraph()
    p_partes.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_partes = p_partes.add_run(f"REQUERENTE: {requerente}\nREQUERIDO: {requerido}")
    format_run(run_partes, font_name='Calibri', size_pt=11, bold=True, color_rgb=(0,0,0))
    
    for _ in range(4):
        doc.add_paragraph()
        
    # Rodapé da Capa
    p_data = doc.add_paragraph()
    p_data.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_data = p_data.add_run(f"Vistoria realizada em: {data_vistoria} às {hora_vistoria}\nLaudo Técnico Final")
    format_run(run_data, font_name='Calibri', size_pt=10, italic=True, color_rgb=(120,120,120))
    
    # Quebra de página para o início do conteúdo
    doc.add_page_break()
    
    # --- CONFIGURAR CABEÇALHO E RODAPÉ DAS PÁGINAS SEGUINTES ---
    section = doc.sections[0]
    # python-docx aplica o cabeçalho em todas as páginas, mas podemos configurar um rodapé numerado simples
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_f = footer_p.add_run(f"Processo nº {processo} | ")
    format_run(run_f, font_name='Calibri', size_pt=9, italic=True, color_rgb=(120,120,120))
    # Adiciona número de página por XML simples do Word
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'PAGE')
    footer_p._p.append(fldSimple)
    
    # --- SEÇÃO 1: APRESENTAÇÃO E OBJETO ---
    h1_1 = doc.add_paragraph()
    run_h1_1 = h1_1.add_run("1. INTRODUÇÃO E DADOS GERAIS")
    format_run(run_h1_1, font_name='Calibri', size_pt=14, bold=True, color_rgb=(47, 84, 150))
    h1_1.paragraph_format.space_before = Pt(12)
    h1_1.paragraph_format.space_after = Pt(6)
    
    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.space_after = Pt(6)
    p_intro.paragraph_format.line_spacing = 1.15
    run_intro = p_intro.add_run(
        f"Este documento apresenta o Laudo de Vistoria de Engenharia Civil elaborado para as partes "
        f"acima identificadas, com o objetivo de registrar, constatar e mapear as anomalias e "
        f"manifestações patológicas presentes no imóvel objeto de perícia perante a Vara da Comarca de {comarca}."
    )
    format_run(run_intro, size_pt=11)
    
    # Tabela de Resumo de Dados de Identificação
    table_id = doc.add_table(rows=5, cols=2)
    table_id.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_id.autofit = False
    
    headers_id = [
        ("Número do Processo", processo),
        ("Ação / Tipo de Perícia", tipo_pericia),
        ("Requerente vs Requerido", f"{requerente} vs {requerido}"),
        ("Endereço Completo", endereco),
        ("Coordenadas Geográficas (Obra)", f"Latitude: {latitude:.6f} | Longitude: {longitude:.6f}")
    ]
    
    for idx, (label, val) in enumerate(headers_id):
        row = table_id.rows[idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        
        # Larguras fixas
        cell_lbl.width = Inches(2.2)
        cell_val.width = Inches(3.8)
        
        set_cell_margins(cell_lbl, top=80, bottom=80, left=100, right=100)
        set_cell_margins(cell_val, top=80, bottom=80, left=100, right=100)
        set_cell_background(cell_lbl, "F2F2F2") # Cinza claro
        
        p_lbl = cell_lbl.paragraphs[0]
        run_lbl = p_lbl.add_run(label)
        format_run(run_lbl, size_pt=10, bold=True, color_rgb=(47, 84, 150))
        
        p_val = cell_val.paragraphs[0]
        run_val = p_val.add_run(str(val))
        format_run(run_val, size_pt=10, color_rgb=(0,0,0))
        
    doc.add_paragraph() # Espaço
    
    # --- SEÇÃO 2: METODOLOGIA E EQUIPAMENTOS ---
    h1_2 = doc.add_paragraph()
    run_h1_2 = h1_2.add_run("2. METODOLOGIA E EQUIPAMENTOS")
    format_run(run_h1_2, font_name='Calibri', size_pt=14, bold=True, color_rgb=(47, 84, 150))
    h1_2.paragraph_format.space_before = Pt(12)
    h1_2.paragraph_format.space_after = Pt(6)
    
    p_met = doc.add_paragraph()
    p_met.paragraph_format.line_spacing = 1.15
    p_met.paragraph_format.space_after = Pt(6)
    run_met = p_met.add_run(
        "A vistoria de campo foi conduzida através de inspeção sensorial detalhada (visuográfica e táctil). "
        "As manifestações patológicas associadas a fissuras, trincas e rachaduras foram quantificadas e qualificadas "
        "com o auxílio de um Fissurômetro de Precisão da marca Trident (comprimento total de 120.0 mm e réguas de "
        "comparação entre 0.05 mm e 1.5 mm). Adicionalmente, as análises de imagem foram submetidas ao módulo auxiliar "
        "de Visão Computacional de alta resolução para validação das espessuras das fendas com base no objeto de escala calibrado, "
        "seguindo as diretrizes técnicas de classificação das normas brasileiras ABNT NBR 9575 (Impermeabilização) e NBR 15575 (Desempenho)."
    )
    format_run(run_met, size_pt=11)
    
    # --- SEÇÃO 3: CONSTATAÇÃO DE DANOS E PATOLOGIAS ---
    h1_3 = doc.add_paragraph()
    run_h1_3 = h1_3.add_run("3. CONSTATAÇÃO DAS PATOLOGIAS (LAUDO FOTOGRÁFICO)")
    format_run(run_h1_3, font_name='Calibri', size_pt=14, bold=True, color_rgb=(47, 84, 150))
    h1_3.paragraph_format.space_before = Pt(12)
    h1_3.paragraph_format.space_after = Pt(12)
    
    if not danos:
        p_no = doc.add_paragraph()
        run_no = p_no.add_run("Nenhuma patologia foi cadastrada nesta vistoria até o momento.")
        format_run(run_no, size_pt=11, italic=True)
    else:
        for idx, d in enumerate(danos, start=1):
            h2 = doc.add_paragraph()
            run_h2 = h2.add_run(f"3.{idx} - {d['ambiente'].upper()} / {d['elemento'].upper()} — {d['local_exato'].upper()}")
            format_run(run_h2, font_name='Calibri', size_pt=12, bold=True, color_rgb=(47, 84, 150))
            h2.paragraph_format.space_before = Pt(10)
            h2.paragraph_format.space_after = Pt(6)
            h2.paragraph_format.keep_with_next = True
            
            # Tabela de especificações técnicas do dano
            table_d = doc.add_table(rows=4, cols=2)
            table_d.alignment = WD_TABLE_ALIGNMENT.CENTER
            table_d.autofit = False
            
            row_data = [
                ("Tipo de Patologia / Anomalia", d['dano']),
                ("Medições e Espessura Calculada", d['medicoes']),
                ("Descrição Técnica Detalhada", d['descricao']),
                ("Observações Complementares", d['observacoes'])
            ]
            
            for d_idx, (lbl, val) in enumerate(row_data):
                row = table_d.rows[d_idx]
                c_lbl, c_val = row.cells[0], row.cells[1]
                
                c_lbl.width = Inches(2.2)
                c_val.width = Inches(3.8)
                
                set_cell_margins(c_lbl, top=80, bottom=80, left=100, right=100)
                set_cell_margins(c_val, top=80, bottom=80, left=100, right=100)
                set_cell_background(c_lbl, "F9FBFD") # leve tom de azul/cinza
                
                p_c_lbl = c_lbl.paragraphs[0]
                run_c_lbl = p_c_lbl.add_run(lbl)
                format_run(run_c_lbl, size_pt=10, bold=True, color_rgb=(47, 84, 150))
                
                p_c_val = c_val.paragraphs[0]
                run_c_val = p_c_val.add_run(str(val))
                format_run(run_c_val, size_pt=10, color_rgb=(0,0,0))
                
            doc.add_paragraph().paragraph_format.space_after = Pt(4) # pequeno espaçamento antes das fotos
            
            # Adiciona as imagens lado a lado
            # As fotos estão no fotos_dir
            # d['fotos'] é uma lista de nomes de arquivos, ex: ['FOTO_001_local.jpg', 'FOTO_001_detalhe.jpg']
            img1 = ""
            img2 = ""
            
            if len(d['fotos']) >= 1:
                img1 = os.path.join(fotos_dir, d['fotos'][0])
            if len(d['fotos']) >= 2:
                img2 = os.path.join(fotos_dir, d['fotos'][1])
                
            cap1 = f"Foto 3.{idx}.a: Visão Geral / Local de inspeção"
            cap2 = f"Foto 3.{idx}.b: Close de Detalhe e Espessura da Patologia"
            
            add_images_side_by_side(doc, img1, img2, cap1, cap2)
            
            p_sep = doc.add_paragraph()
            p_sep.paragraph_format.space_after = Pt(12)
            p_sep.add_run("—" * 65).font.color.rgb = RGBColor(220, 220, 220)
            
    # --- SEÇÃO 4: CONCLUSÃO E ASSINATURA ---
    doc.add_page_break()
    
    h1_4 = doc.add_paragraph()
    run_h1_4 = h1_4.add_run("4. CONSIDERAÇÕES FINAIS E ENCERRAMENTO")
    format_run(run_h1_4, font_name='Calibri', size_pt=14, bold=True, color_rgb=(47, 84, 150))
    h1_4.paragraph_format.space_before = Pt(12)
    h1_4.paragraph_format.space_after = Pt(6)
    
    p_concl = doc.add_paragraph()
    p_concl.paragraph_format.line_spacing = 1.15
    p_concl.paragraph_format.space_after = Pt(12)
    run_concl = p_concl.add_run(
        "As anomalias e anotações técnicas foram descritas e registradas em conformidade com as boas práticas de engenharia diagnóstica, "
        "utilizando equipamentos devidamente calibrados e instrumentação técnica recomendada para laudos judiciais. "
        "Apresenta-se o presente laudo em fé de seu fiel cumprimento, contendo o registro fidedigno do estado atual da obra, "
        "restando o profissional perito à inteira disposição deste Douto Juízo para prestação de esclarecimentos cabíveis."
    )
    format_run(run_concl, size_pt=11)
    
    for _ in range(4):
        doc.add_paragraph()
        
    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sig.paragraph_format.keep_with_next = True
    run_sig_line = p_sig.add_run("_____________________________________________________\n")
    format_run(run_sig_line, size_pt=11, bold=True, color_rgb=(100, 100, 100))
    
    run_sig_title = p_sig.add_run("PERITO DE ENGENHARIA CIVIL DESIGNADO\nCREA/CONFEA SP — Perito Judicial")
    format_run(run_sig_title, size_pt=11, bold=True, color_rgb=(47, 84, 150))
    
    # Salva o Laudo em formato .docx
    doc.save(output_docx_path)
    print(f"Laudo gerado com sucesso em: {output_docx_path}")

if __name__ == "__main__":
    # Exemplo de teste standalone
    import sys
    script_dir = os.path.dirname(os.path.abspath(__file__))
    excel = os.path.join(script_dir, "formulario_vistoria.xlsx")
    fotos = os.path.join(script_dir, "fotos")
    out = os.path.join(script_dir, "laudo_pericial_teste.docx")
    
    if os.path.exists(excel):
        generate_docx_laudo(excel, fotos, out)
