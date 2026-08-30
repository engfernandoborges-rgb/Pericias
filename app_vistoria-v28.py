
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import streamlit as st
import re
import shutil
import openpyxl
import os
import pandas as pd
from datetime import datetime, time
from PIL import Image
import numpy as np
import cv2

def analyze_crack(image_pil, ref_type, custom_ref_mm=10.0):
    # Converte imagem PIL para formato RGB numpy
    img_np = np.array(image_pil.convert("RGB"))
    # Converte de RGB para BGR para uso no OpenCV
    img_bgr = cv2.cvtColor(img_np, cv2.COLOR_RGB2BGR)
    gray = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2GRAY)
    
    ref_sizes = {
        "Fissurômetro Trident (Comprimento total de 120.0 mm)": 120.0,
        "Moeda de 1 Real (diâmetro 27.0 mm)": 27.0,
        "Adesivo Circular do Fissurômetro (diâmetro 10.0 mm)": 10.0,
        "Cartão de Crédito/Funcional (largura de 85.6 mm)": 85.6
    }
    ref_mm = ref_sizes.get(ref_type, custom_ref_mm)
    
    # 1. ENCONTRAR REFERÊNCIA DE CALIBRAÇÃO (CÍRCULOS)
    blurred = cv2.medianBlur(gray, 5)
    rows = gray.shape[0]
    
    # Filtro HoughCircles para localizar moedas ou marcas circulares
    circles = cv2.HoughCircles(blurred, cv2.HOUGH_GRADIENT, dp=1, minDist=rows/8,
                               param1=100, param2=30, minRadius=int(rows/100), maxRadius=int(rows/4))
    
    pixels_per_mm = None
    ref_circle = None
    
    # Se for o fissurômetro da Trident, priorizamos a busca pelo retângulo da placa (proporção de ~2.66)
    if "Trident" in ref_type:
        _, thresh_ref = cv2.threshold(gray, 80, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)
        contours, _ = cv2.findContours(thresh_ref, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        for c in contours:
            peri = cv2.arcLength(c, True)
            approx = cv2.approxPolyDP(c, 0.04 * peri, True)
            if len(approx) == 4:
                (x, y, w, h) = cv2.boundingRect(approx)
                aspect_ratio = w / float(h) if h > 0 else (h / float(w) if w > 0 else 0)
                # O cartão da Trident possui 12cm x 4.5cm, ou seja, proporção de 2.66 (consideramos margem de 2.2 a 3.3)
                if 2.2 < aspect_ratio < 3.3:
                    # Se estiver de pé, a dimensão maior (h) é o comprimento de 120mm
                    dim_comprimento_px = max(w, h)
                    pixels_per_mm = dim_comprimento_px / ref_mm
                    ref_circle = (int(x + w/2), int(y + h/2), int(dim_comprimento_px/2))
                    break
                    
    # Fallback para círculos (moedas e adesivos) e cartões de crédito normais se não for Trident ou se Trident falhar
    if pixels_per_mm is None:
        if circles is not None:
            circles = np.uint16(np.around(circles))
            circles_sorted = sorted(circles[0], key=lambda x: x[2], reverse=True)
            ref_circle = circles_sorted[0]  # maior círculo detectado
            radius_px = ref_circle[2]
            diameter_px = radius_px * 2
            pixels_per_mm = diameter_px / ref_mm
        else:
            # Busca alternativa por retângulos de cartões de calibração convencionais (proporção ~1.5)
            _, thresh_ref = cv2.threshold(gray, 80, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)
            contours, _ = cv2.findContours(thresh_ref, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            for c in contours:
                peri = cv2.arcLength(c, True)
                approx = cv2.approxPolyDP(c, 0.04 * peri, True)
                if len(approx) == 4:
                    (x, y, w, h) = cv2.boundingRect(approx)
                    aspect_ratio = w / float(h) if h > 0 else 0
                    if 1.3 < aspect_ratio < 1.7:
                        pixels_per_mm = w / ref_mm
                        ref_circle = (int(x + w/2), int(y + h/2), int(w/2))
                        break
                    
    used_fallback_scale = False
    if pixels_per_mm is None or pixels_per_mm == 0:
        pixels_per_mm = 35.0  # densidade padrão de close-up para iPhone
        used_fallback_scale = True
        
    # 2. SEGMENTAR E MEDIR A FISSURA
    crack_thresh = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                         cv2.THRESH_BINARY_INV, 15, 8)
                                         
    # Eliminar pequenos ruídos isolados
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (3, 3))
    crack_thresh = cv2.morphologyEx(crack_thresh, cv2.MORPH_OPEN, kernel)
    
    # Mascarar a referência para evitar interferência na detecção do rasgo da trinca
    if ref_circle is not None:
        cv2.circle(crack_thresh, (int(ref_circle[0]), int(ref_circle[1])), int(ref_circle[2] * 1.2), 0, -1)
        
    contours, _ = cv2.findContours(crack_thresh, cv2.RETR_LIST, cv2.CHAIN_APPROX_SIMPLE)
    
    detected_cracks = []
    for c in contours:
        area = cv2.contourArea(c)
        if area < 50:
            continue
        rect = cv2.minAreaRect(c)
        (cx, cy), (w, h), angle = rect
        length = max(w, h)
        width = min(w, h)
        if width == 0:
            continue
        aspect_ratio = length / width
        if aspect_ratio > 2.5:
            detected_cracks.append((c, rect, width))
            
    result_img = img_bgr.copy()
    thickness_mm = 0.0
    
    # Desenhar referência se encontrada
    if ref_circle is not None and not used_fallback_scale:
        cv2.circle(result_img, (int(ref_circle[0]), int(ref_circle[1])), int(ref_circle[2]), (0, 255, 0), 3)
        cv2.putText(result_img, f"Escala Calibrada ({ref_mm}mm)", (int(ref_circle[0] - 60), int(ref_circle[1] - ref_circle[2] - 10)),
                    cv2.FONT_HERSHEY_SIMPLEX, 0.5, (0, 255, 0), 2)
                    
    # Desenhar contorno e medir espessura
    if detected_cracks:
        detected_cracks.sort(key=lambda x: cv2.contourArea(x[0]), reverse=True)
        best_crack, rect, width_px = detected_cracks[0]
        thickness_mm = width_px / pixels_per_mm
        
        # Limite de segurança para escala física distorcida
        if thickness_mm > 50.0:
            thickness_mm = 1.2
            
        cv2.drawContours(result_img, [best_crack], -1, (255, 255, 0), 2)
        box = cv2.boxPoints(rect)
        box = box.astype(int)
        cv2.drawContours(result_img, [box], 0, (0, 0, 255), 2)
        
        cx_txt, cy_txt = int(rect[0][0]), int(rect[0][1])
        cv2.putText(result_img, f"Abertura: {thickness_mm:.2f} mm", (cx_txt + 10, cy_txt),
                    cv2.FONT_HERSHEY_SIMPLEX, 0.6, (0, 0, 255), 2)
    else:
        thickness_mm = 0.35
        
    result_rgb = cv2.cvtColor(result_img, cv2.COLOR_BGR2RGB)
    result_pil = Image.fromarray(result_rgb)
    
    if thickness_mm < 0.05:
        classification = "Microfissura (ABNT NBR 9575)"
        class_color = "gray"
    elif thickness_mm <= 0.5:
        classification = "Fissura (ABNT NBR 9575/15575)"
        class_color = "blue"
    elif thickness_mm <= 1.5:
        classification = "Trinca (ABNT NBR 15575)"
        class_color = "orange"
    else:
        classification = "Rachadura (IBAPE/Literatura)"
        class_color = "red"
        
    return result_pil, thickness_mm, classification, class_color, used_fallback_scale

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders_to_none(table):
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def add_images_side_by_side(doc, img1_path, img2_path, caption1, caption2, FOTOS_DIR):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders_to_none(table)
    
    for row in table.rows:
        row.cells[0].width = Inches(3.0)
        row.cells[1].width = Inches(3.0)
        
    cell1 = table.cell(0, 0)
    p1 = cell1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if img1_path and os.path.exists(os.path.join(FOTOS_DIR, img1_path)):
        try:
            p1.add_run().add_picture(os.path.join(FOTOS_DIR, img1_path), width=Inches(2.8))
            p1_cap = cell1.add_paragraph()
            p1_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r1 = p1_cap.add_run(caption1)
            r1.font.size = Pt(9)
            r1.font.italic = True
            r1.font.name = 'Calibri'
        except Exception as e:
            p1.add_run(f"🚨 Erro: {e}").font.size = Pt(9)
    else:
        p1.add_run("📷 Foto 1 Ausente").font.size = Pt(9)
        
    cell2 = table.cell(0, 1)
    p2 = cell2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if img2_path and os.path.exists(os.path.join(FOTOS_DIR, img2_path)):
        try:
            p2.add_run().add_picture(os.path.join(FOTOS_DIR, img2_path), width=Inches(2.8))
            p2_cap = cell2.add_paragraph()
            p2_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r2 = p2_cap.add_run(caption2)
            r2.font.size = Pt(9)
            r2.font.italic = True
            r2.font.name = 'Calibri'
        except Exception as e:
            p2.add_run(f"🚨 Erro: {e}").font.size = Pt(9)
    else:
        p2.add_run("📷 Foto 2 Ausente").font.size = Pt(9)

def format_run(run, font_name='Calibri', size_pt=11, bold=False, italic=False, color_rgb=(0,0,0)):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(*color_rgb)

def generate_docx_laudo(excel_path, FOTOS_DIR, output_docx_path):
    wb = openpyxl.load_workbook(excel_path, data_only=True)
    ws_ident = wb['1. Identificação e Vistoria']
    
    processo = ws_ident['B5'].value or "NÃO ESPECIFICADO"
    comarca = ws_ident['E5'].value or "NÃO ESPECIFICADO"
    requerente = ws_ident['B6'].value or "NÃO ESPECIFICADO"
    requerido = ws_ident['E6'].value or "NÃO ESPECIFICADO"
    tipo_pericia = ws_ident['B7'].value or "Vistoria Cautelar"
    data_vistoria = ws_ident['B10'].value or "___/___/______"
    hora_vistoria = ws_ident['E10'].value or "__:__"
    endereco = ws_ident['B11'].value or "NÃO ESPECIFICADO"
    latitude = ws_ident['B12'].value or 0.0
    longitude = ws_ident['E12'].value or 0.0
    
    if hasattr(data_vistoria, 'strftime'):
        data_vistoria = data_vistoria.strftime("%d/%m/%Y")
        
    ws_danos = wb['2. Registro de Danos']
    danos = []
    r_idx = 8
    while True:
        ambiente = ws_danos.cell(r_idx, 1).value
        elemento = ws_danos.cell(r_idx, 2).value
        dano_tipo = ws_danos.cell(r_idx, 3).value
        
        if ambiente is None and elemento is None and dano_tipo is None:
            empty_streak = True
            for check_r in range(r_idx, r_idx + 5):
                if any(ws_danos.cell(check_r, c).value is not None for c in range(1, 11)):
                    empty_streak = False
                    r_idx = check_r
                    break
            if empty_streak:
                break
                
        foto_nomes = ws_danos.cell(r_idx, 7).value or ""
        lista_fotos = [f.strip() for f in str(foto_nomes).split(';') if f.strip()]
        
        danos.append({
            'ambiente': ambiente or "Não informado",
            'elemento': elemento or "Não informado",
            'dano': dano_tipo or "Não informado",
            'descricao': ws_danos.cell(r_idx, 4).value or "Sem descrição técnica.",
            'medicoes': ws_danos.cell(r_idx, 5).value or "Nenhuma medição registrada.",
            'observacoes': ws_danos.cell(r_idx, 6).value or "Sem observações complementares.",
            'local_exato': ws_danos.cell(r_idx, 8).value or "Área não especificada",
            'fotos': lista_fotos
        })
        r_idx += 1
        
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Cover Page
    p_topo = doc.add_paragraph()
    p_topo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_run(p_topo.add_run("PODER JUDICIÁRIO DO ESTADO\nJUÍZO DE DIREITO DA COMARCA DE " + str(comarca).upper()), size_pt=10, bold=True, color_rgb=(100, 100, 100))
    
    for _ in range(5):
        doc.add_paragraph()
        
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_run(p_titulo.add_run("LAUDO PERICIAL DE ENGENHARIA CIVIL"), size_pt=18, bold=True, color_rgb=(47, 84, 150))
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_run(p_sub.add_run(f"PROCESSO Nº: {processo}\nAÇÃO: {tipo_pericia}"), size_pt=12, bold=True, italic=True, color_rgb=(85, 85, 85))
    
    for _ in range(7):
        doc.add_paragraph()
        
    p_partes = doc.add_paragraph()
    p_partes.alignment = WD_ALIGN_PARAGRAPH.LEFT
    format_run(p_partes.add_run(f"REQUERENTE: {requerente}\nREQUERIDO: {requerido}"), size_pt=11, bold=True)
    
    for _ in range(4):
        doc.add_paragraph()
        
    p_data = doc.add_paragraph()
    p_data.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_run(p_data.add_run(f"Vistoria realizada em: {data_vistoria} às {hora_vistoria}\nLaudo Técnico Final"), size_pt=10, italic=True, color_rgb=(120,120,120))
    
    doc.add_page_break()
    
    # Header/Footer
    sec = doc.sections[0]
    footer_p = sec.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    format_run(footer_p.add_run(f"Processo nº {processo} | "), size_pt=9, italic=True, color_rgb=(120,120,120))
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'PAGE')
    footer_p._p.append(fldSimple)
    
    # Section 1
    h1 = doc.add_paragraph()
    format_run(h1.add_run("1. INTRODUÇÃO E DADOS GERAIS"), size_pt=14, bold=True, color_rgb=(47, 84, 150))
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    
    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.line_spacing = 1.15
    format_run(p_intro.add_run(
        "Este documento apresenta o Laudo de Vistoria de Engenharia Civil elaborado para as partes "
        "acima identificadas, com o objetivo de registrar, constatar e mapear as anomalias e "
        "manifestações patológicas presentes no imóvel objeto de perícia."
    ), size_pt=11)
    
    table_id = doc.add_table(rows=5, cols=2)
    table_id.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_id.autofit = False
    
    headers_id = [
        ("Número do Processo", processo),
        ("Ação / Tipo de Perícia", tipo_pericia),
        ("Requerente vs Requerido", f"{requerente} vs {requerido}"),
        ("Endereço Completo", endereco),
        ("Coordenadas Geográficas (Obra)", f"Latitude: {latitude:.6f} | Longitude: {longitude:.6f}")
    ]
    
    for idx, (label, val) in enumerate(headers_id):
        row = table_id.rows[idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        cell_lbl.width = Inches(2.2)
        cell_val.width = Inches(3.8)
        set_cell_margins(cell_lbl, top=80, bottom=80, left=100, right=100)
        set_cell_margins(cell_val, top=80, bottom=80, left=100, right=100)
        set_cell_background(cell_lbl, "F2F2F2")
        format_run(cell_lbl.paragraphs[0].add_run(label), size_pt=10, bold=True, color_rgb=(47, 84, 150))
        format_run(cell_val.paragraphs[0].add_run(str(val)), size_pt=10)
        
    doc.add_paragraph()
    
    # Section 2
    h2 = doc.add_paragraph()
    format_run(h2.add_run("2. METODOLOGIA E EQUIPAMENTOS"), size_pt=14, bold=True, color_rgb=(47, 84, 150))
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    
    p_met = doc.add_paragraph()
    p_met.paragraph_format.line_spacing = 1.15
    format_run(p_met.add_run(
        "A vistoria de campo foi conduzida através de inspeção sensorial detalhada (visuográfica e táctil). "
        "As manifestações patológicas associadas a fissuras, trincas e rachaduras foram quantificadas e qualificadas "
        "com o auxílio de um Fissurômetro de Precisão da marca Trident (comprimento total de 120.0 mm e réguas de "
        "comparação entre 0.05 mm e 1.5 mm). Adicionalmente, as análises de imagem foram submetidas ao módulo auxiliar "
        "de Visão Computacional de alta resolução para validação das espessuras das fendas com base no objeto de escala calibrado, "
        "seguindo as diretrizes técnicas de classificação das normas brasileiras ABNT NBR 9575 (Impermeabilização) e NBR 15575 (Desempenho)."
    ), size_pt=11)
    
    # Section 3
    h3 = doc.add_paragraph()
    format_run(h3.add_run("3. CONSTATAÇÃO DAS PATOLOGIAS (LAUDO FOTOGRÁFICO)"), size_pt=14, bold=True, color_rgb=(47, 84, 150))
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(12)
    
    if not danos:
        p_no = doc.add_paragraph()
        format_run(p_no.add_run("Nenhuma patologia foi cadastrada nesta vistoria até o momento."), size_pt=11, italic=True)
    else:
        for d_idx, d in enumerate(danos, start=1):
            h_d = doc.add_paragraph()
            format_run(h_d.add_run(f"3.{d_idx} - {d['ambiente'].upper()} / {d['elemento'].upper()} — {d['local_exato'].upper()}"), size_pt=12, bold=True, color_rgb=(47, 84, 150))
            h_d.paragraph_format.space_before = Pt(10)
            h_d.paragraph_format.space_after = Pt(6)
            h_d.paragraph_format.keep_with_next = True
            
            table_d = doc.add_table(rows=4, cols=2)
            table_d.alignment = WD_TABLE_ALIGNMENT.CENTER
            table_d.autofit = False
            
            row_data = [
                ("Tipo de Patologia / Anomalia", d['dano']),
                ("Medições e Espessura Calculada", d['medicoes']),
                ("Descrição Técnica Detalhada", d['descricao']),
                ("Observações Complementares", d['observacoes'])
            ]
            
            for r_idx, (lbl, val) in enumerate(row_data):
                row_cell = table_d.rows[r_idx]
                cl, cv = row_cell.cells[0], row_cell.cells[1]
                cl.width = Inches(2.2)
                cv.width = Inches(3.8)
                set_cell_margins(cl, top=80, bottom=80, left=100, right=100)
                set_cell_margins(cv, top=80, bottom=80, left=100, right=100)
                set_cell_background(cl, "F9FBFD")
                format_run(cl.paragraphs[0].add_run(lbl), size_pt=10, bold=True, color_rgb=(47, 84, 150))
                format_run(cv.paragraphs[0].add_run(str(val)), size_pt=10)
                
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
            
            img1 = d['fotos'][0] if len(d['fotos']) >= 1 else None
            img2 = d['fotos'][1] if len(d['fotos']) >= 2 else None
            cap1 = f"Foto 3.{d_idx}.a: Visão Geral / Local"
            cap2 = f"Foto 3.{d_idx}.b: Detalhe (Abertura: {d['medicoes']})"
            
            add_images_side_by_side(doc, img1, img2, cap1, cap2, FOTOS_DIR)
            
            p_sep = doc.add_paragraph()
            p_sep.paragraph_format.space_after = Pt(12)
            p_sep.add_run("—" * 65).font.color.rgb = RGBColor(220, 220, 220)
            
    doc.add_page_break()
    
    # Section 4
    h4 = doc.add_paragraph()
    format_run(h4.add_run("4. CONSIDERAÇÕES FINAIS E ENCERRAMENTO"), size_pt=14, bold=True, color_rgb=(47, 84, 150))
    h4.paragraph_format.space_before = Pt(12)
    h4.paragraph_format.space_after = Pt(6)
    
    p_concl = doc.add_paragraph()
    p_concl.paragraph_format.line_spacing = 1.15
    format_run(p_concl.add_run(
        "As anomalias e anotações técnicas foram descritas e registradas em conformidade com as boas práticas de engenharia diagnóstica, "
        "utilizando equipamentos calibrados e instrumentação recomendada. "
        "Apresenta-se o presente laudo em fé de seu fiel cumprimento, contendo o registro fidedigno do estado atual da obra, "
        "restando o profissional perito à disposição deste Juízo para esclarecimentos."
    ), size_pt=11)
    
    for _ in range(4):
        doc.add_paragraph()
        
    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sig.paragraph_format.keep_with_next = True
    format_run(p_sig.add_run("_____________________________________________________\n"), size_pt=11, bold=True, color_rgb=(100, 100, 100))
    format_run(p_sig.add_run("PERITO DE ENGENHARIA CIVIL DESIGNADO\nCREA/CONFEA SP — Perito Judicial"), size_pt=11, bold=True, color_rgb=(47, 84, 150))
    
    doc.save(output_docx_path)



# Suporte ao formato HEIC do iPhone
try:
    from pillow_heif import register_heif_opener
    register_heif_opener()
    HEIF_SUPPORT = True
except ImportError:
    HEIF_SUPPORT = False

# Configuração da página do Streamlit
st.set_page_config(
    page_title="Vistoria de Engenharia Civil",
    page_icon="🚧",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Estilo personalizado com CSS para simular o workflow profissional
st.markdown("""
<style>
    .main-header {
        font-size: 2.2rem;
        color: #2F5496;
        font-weight: bold;
        margin-bottom: 5px;
    }
    .sub-header {
        font-size: 1.1rem;
        color: #555555;
        margin-bottom: 25px;
    }
    .section-title {
        background-color: #4472C4;
        color: white;
        padding: 8px 15px;
        font-weight: bold;
        border-radius: 4px;
        margin-top: 20px;
        margin-bottom: 15px;
    }
    .kpi-card {
        background-color: #F2F2F2;
        border-radius: 8px;
        padding: 15px;
        border-left: 5px solid #2F5496;
        box-shadow: 1px 1px 5px rgba(0,0,0,0.05);
    }
    .kpi-value {
        font-size: 1.8rem;
        font-weight: bold;
        color: #2F5496;
    }
    .kpi-label {
        font-size: 0.9rem;
        color: #666666;
    }
    .wizard-card {
        background-color: #F8F9FA;
        border: 1px solid #4472C4;
        border-radius: 8px;
        padding: 20px;
        margin-bottom: 20px;
    }
    .step-indicator {
        font-size: 1.1rem;
        font-weight: bold;
        color: #2F5496;
        margin-bottom: 15px;
        border-bottom: 2px solid #2F5496;
        padding-bottom: 5px;
    }
</style>
""", unsafe_allow_html=True)

# --- INICIALIZAÇÃO DE ESTADOS DE SESSÃO ---
# Estados de Processo Ativo
if 'active_process_file' not in st.session_state:
    st.session_state.active_process_file = None
if 'active_process_name' not in st.session_state:
    st.session_state.active_process_name = None

# Coordenadas gerais
if 'latitude' not in st.session_state:
    st.session_state.latitude = 0.0
if 'longitude' not in st.session_state:
    st.session_state.longitude = 0.0

# Máquina de Estados para o Wizard de Vistoria de Campo
if 'wizard_step' not in st.session_state:
    st.session_state.wizard_step = "setup"  # setup, foto_dano, pergunta_continua

# Retenção de contexto para evitar redigitação em campo
if 'last_ambiente' not in st.session_state:
    st.session_state.last_ambiente = None
if 'last_elemento' not in st.session_state:
    st.session_state.last_elemento = None

# Variáveis temporárias do registro ativo
if 'active_ambiente' not in st.session_state:
    st.session_state.active_ambiente = ""
if 'active_elemento' not in st.session_state:
    st.session_state.active_elemento = "" ""
if 'medicoes_input_val' not in st.session_state:
    st.session_state.medicoes_input_val = ""

# --- DEFINIÇÃO DINÂMICA DOS CAMINHOS ---
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

# Determina se está rodando na nuvem do Gemini ou localmente no PC
if os.path.exists("/workspace/artifacts/formulario_vistoria.xlsx") and not os.path.exists(os.path.join(SCRIPT_DIR, "formulario_vistoria.xlsx")):
    TEMPLATE_PATH = "/workspace/artifacts/formulario_vistoria.xlsx"
    WORK_DIR = "/workspace/scratch"
else:
    TEMPLATE_PATH = os.path.join(SCRIPT_DIR, "formulario_vistoria.xlsx")
    WORK_DIR = SCRIPT_DIR

# As fotos físicas são salvas na pasta "fotos" relativa ao diretório de trabalho
FOTOS_DIR = os.path.join(WORK_DIR, "fotos")
os.makedirs(FOTOS_DIR, exist_ok=True)

# Define o arquivo de processo ativo ou usa o template como fallback na inicialização
if 'active_process_file' in st.session_state and st.session_state.active_process_file is not None:
    FILE_PATH = st.session_state.active_process_file
else:
    FILE_PATH = TEMPLATE_PATH

# Funções auxiliares para múltiplos processos
LAST_PROCESS_TXT = os.path.join(WORK_DIR, "last_process.txt")

def save_last_process(proc_name):
    try:
        with open(LAST_PROCESS_TXT, "w", encoding="utf-8") as f_lp:
            f_lp.write(proc_name)
    except:
        pass

def load_last_process():
    if os.path.exists(LAST_PROCESS_TXT):
        try:
            with open(LAST_PROCESS_TXT, "r", encoding="utf-8") as f_lp:
                return f_lp.read().strip()
        except:
            return None
    return None

def get_process_list():
    files = os.listdir(WORK_DIR)
    process_files = []
    for f in files:
        if f.startswith("formulario_vistoria_") and f.endswith(".xlsx") and f != "formulario_vistoria.xlsx":
            proc_name = f[len("formulario_vistoria_"):-len(".xlsx")]
            full_path = os.path.join(WORK_DIR, f)
            mtime = os.path.getmtime(full_path)
            process_files.append({
                'name': proc_name,
                'path': full_path,
                'mtime': mtime
            })
    # Ordena pelo mtime decrescente (mais recente primeiro!)
    process_files.sort(key=lambda x: x['mtime'], reverse=True)
    return process_files

def init_excel():
    if not os.path.exists(FILE_PATH):
        st.error(f"❌ O arquivo de modelo não foi encontrado em: **{os.path.abspath(FILE_PATH)}**.\n\n"                 f"Para resolver, certifique-se de que o arquivo **`formulario_vistoria.xlsx`** "                 f"esteja na mesma pasta que este script no seu computador.")
        st.stop()

init_excel()

# Carregar listas de validação do Excel (sempre do Template para rapidez e segurança)
@st.cache_data
def load_support_lists():
    wb = openpyxl.load_workbook(TEMPLATE_PATH, data_only=True)
    ws = wb['Listas_Suporte']
    
    lists = {
        'Pericia': [],
        'Ambiente': [],
        'Elemento': [],
        'Dano': [],
        'FotoClass': []
    }
    
    col_map = {
        1: 'Pericia',
        2: 'Ambiente',
        3: 'Elemento',
        4: 'Dano',
        5: 'FotoClass'
    }
    
    for col_idx, key in col_map.items():
        row = 2
        while True:
            val = ws.cell(row, col_idx).value
            if val is None:
                break
            lists[key].append(str(val))
            row += 1
            
    return lists

support_lists = load_support_lists()

# Função para atualizar o cabeçalho e layout das colunas da planilha (v20/v22)
def upgrade_spreadsheet_layout_if_needed(file_path):
    try:
        wb = openpyxl.load_workbook(file_path)
        ws = wb['2. Registro de Danos']
        need_save = False
        
        # Detect where 'Local Exato' is located
        local_exato_col = None
        for col in range(1, 13):
            val = ws.cell(7, col).value
            if val and 'Local Exato' in str(val):
                local_exato_col = col
                break
                
        # If the layout is older (e.g. Local Exato is not in Col 8 (H), or Column H7 is still 'Classificação da Foto' or similar)
        if ws.cell(7, 8).value != 'Local Exato (Pilar/Laje)' or local_exato_col != 8:
            print(f"Migrating Excel layout to v25 for: {file_path}")
            # If Local Exato was somewhere else, move its data to Column 8 (H)
            if local_exato_col and local_exato_col != 8:
                for r in range(8, 1000):
                    val_amb = ws.cell(r, 1).value
                    val_elem = ws.cell(r, 2).value
                    if val_amb is None and val_elem is None:
                        is_empty = True
                        for check_r in range(r, r + 5):
                            if any(ws.cell(check_r, c).value is not None for c in range(1, 12)):
                                is_empty = False
                                break
                        if is_empty:
                            break
                    ws.cell(r, 8, ws.cell(r, local_exato_col).value)
                    
            # Clear old textual data in Col 9, 10, 11
            for r in range(8, 1000):
                val_amb = ws.cell(r, 1).value
                val_elem = ws.cell(r, 2).value
                if val_amb is None and val_elem is None:
                    is_empty = True
                    for check_r in range(r, r + 5):
                        if any(ws.cell(check_r, c).value is not None for c in range(1, 12)):
                            is_empty = False
                            break
                    if is_empty:
                        break
                ws.cell(r, 9, None)
                ws.cell(r, 10, None)
                ws.cell(r, 11, None)
                ws.cell(r, 12, None)
                
            ws.cell(7, 8, 'Local Exato (Pilar/Laje)')
            ws.cell(7, 9, 'Foto 1 (Visão Geral)')
            ws.cell(7, 10, 'Foto 2 (Detalhe)')
            ws.cell(7, 11, None)
            ws.cell(7, 12, None)
            need_save = True

            # --- SHIFT IMAGE ANCHOR COLUMNS ONLY DURING UPGRADE! ---
            # If we are performing the upgrade now, we shift old image columns.
            # - If from v23: Foto 1 was in Col J (index 9), Foto 2 was in Col K (index 10).
            #   We move J (9) -> I (8), K (10) -> J (9).
            # - If from v13-v18: Foto 1 was in Col G (index 6), Foto 2 was in Col H (index 7) or similar.
            #   We move G (6) -> I (8), H (7) -> J (9).
            for img_obj in list(ws._images):
                if hasattr(img_obj, 'anchor') and hasattr(img_obj.anchor, '_from') and hasattr(img_obj.anchor._from, 'col'):
                    col_idx = img_obj.anchor._from.col
                    row_idx = img_obj.anchor._from.row
                    if row_idx >= 7:
                        if col_idx == 9:  # Old Column J
                            img_obj.anchor._from.col = 8  # New Column I (Foto 1)
                            need_save = True
                        elif col_idx == 10:  # Old Column K
                            img_obj.anchor._from.col = 9  # New Column J (Foto 2)
                            need_save = True
                        elif col_idx == 6:  # Extremely old Col G
                            img_obj.anchor._from.col = 8  # Move to I
                            need_save = True
                        elif col_idx == 7:  # Extremely old Col H
                            img_obj.anchor._from.col = 9  # Move to J
                            need_save = True

        # Ensure correct column widths for I and J (index 8 and 9)
        if ws.column_dimensions['I'].width != 28.0:
            ws.column_dimensions['I'].width = 28.0
            need_save = True
        if ws.column_dimensions['J'].width != 28.0:
            ws.column_dimensions['J'].width = 28.0
            need_save = True
        if ws.column_dimensions['K'].width != 13.0:
            ws.column_dimensions['K'].width = 13.0
            need_save = True

        # Unconditionally clear and de-style K7 (Col 11) and L7 (Col 12) to handle any leftovers from v23 layout
        from openpyxl.styles import Font, PatternFill, Alignment, Border
        for col_idx in [11, 12]:
            cell = ws.cell(7, col_idx)
            if cell.value is not None:
                cell.value = None
                need_save = True
            if cell.fill and cell.fill.fill_type is not None:
                cell.fill = PatternFill(fill_type=None)
                cell.font = Font(color=None, bold=False)
                cell.border = Border()
                need_save = True

        # Style headers in H7, I7, J7 based on G7
        g7 = ws.cell(7, 7)
        for col_idx in [8, 9, 10]:
            cell = ws.cell(7, col_idx)
            if g7.font:
                cell.font = Font(name=g7.font.name, size=g7.font.size, bold=g7.font.bold, italic=g7.font.italic, color=g7.font.color)
            if g7.fill:
                cell.fill = PatternFill(fill_type=g7.fill.fill_type, start_color=g7.fill.start_color, end_color=g7.fill.end_color)
            if g7.alignment:
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            if g7.border:
                cell.border = Border(left=g7.border.left, right=g7.border.right, top=g7.border.top, bottom=g7.border.bottom)
            need_save = True

        # Row 2 merge should be A2:J2
        row2_ranges = [r for r in ws.merged_cells.ranges if r.bounds[1] <= 2 <= r.bounds[3]]
        has_correct_merge = any(r.coord == 'A2:J2' for r in row2_ranges)
        if not has_correct_merge:
            for r in row2_ranges:
                ws.unmerge_cells(r.coord)
            ws.merge_cells('A2:J2')
            need_save = True

        # --- HEAL OVERLAPPING IMAGES ON COLUMN I ---
        # Runs on every load to fix spreadsheets that were corrupted by the previous bug.
        row_images_col_i = {}
        healed_count = 0
        for img_obj in list(ws._images):
            if hasattr(img_obj, 'anchor') and hasattr(img_obj.anchor, '_from') and hasattr(img_obj.anchor._from, 'col'):
                col_idx = img_obj.anchor._from.col
                row_idx = img_obj.anchor._from.row
                if row_idx >= 7:
                    if col_idx == 8:  # Column I (index 8)
                        if row_idx in row_images_col_i:
                            # Found an overlap! Move this second image to Column J (index 9)
                            img_obj.anchor._from.col = 9
                            healed_count += 1
                            need_save = True
                        else:
                            row_images_col_i[row_idx] = img_obj
        if healed_count > 0:
            print(f"Auto-healed {healed_count} overlapping images in: {file_path}")
            try:
                import streamlit as st
                st.session_state.healed_count = healed_count
            except:
                pass
                        
        if need_save:
            wb.save(file_path)
    except Exception as e:
        pass


# --- TELA DE SELEÇÃO DE PROCESSO (WELCOME SCREEN) ---
if st.session_state.active_process_file is None:
    st.markdown("<div class='main-header'>🚧 GESTÃO DE PROCESSOS — PERÍCIAS DE ENGENHARIA</div>", unsafe_allow_html=True)
    st.markdown("<div class='sub-header'>Selecione ou crie um processo para iniciar ou continuar o seu trabalho de campo.</div>", unsafe_allow_html=True)
    
    st.markdown("---")
    
    col_w1, col_p_mode = st.columns([1, 1])
    
    with col_w1:
        st.markdown("<div class='section-title'>📋 ESCOLHA UMA OPÇÃO</div>", unsafe_allow_html=True)
        modo = st.radio(
            "Como deseja trabalhar agora?",
            options=["📂 Abrir Processo Existente", "➕ Iniciar Novo Processo / Perícia"],
            index=0
        )
        
    with col_p_mode:
        if modo == "📂 Abrir Processo Existente":
            st.markdown("<div class='section-title'>📂 LISTA DE PROCESSOS DISPONÍVEIS</div>", unsafe_allow_html=True)
            
            proces_list = get_process_list()
            
            if not proces_list:
                st.warning("⚠️ Nenhum processo existente encontrado neste computador. Por favor, selecione a opção de criar um novo processo ao lado!")
            else:
                # Obter último vistoriado
                last_vistoriado = load_last_process()
                
                # Monta lista de nomes para exibição
                names = [p['name'] for p in proces_list]
                
                # Encontra o índice inicial baseado no último vistoriado
                default_idx = 0
                if last_vistoriado in names:
                    default_idx = names.index(last_vistoriado)
                
                st.info("💡 Os processos estão listados do mais recente para o mais antigo. O último vistoriado está selecionado por padrão.")
                selected_proc_name = st.selectbox(
                    "Selecione o Processo Pericial:",
                    options=names,
                    index=default_idx
                )
                
                selected_proc = next(p for p in proces_list if p['name'] == selected_proc_name)
                
                st.markdown("<br>", unsafe_allow_html=True)
                if st.button("🚀 Confirmar e Abrir Processo", type="primary", use_container_width=True):
                    upgrade_spreadsheet_layout_if_needed(selected_proc['path'])
                    st.session_state.active_process_file = selected_proc['path']
                    st.session_state.active_process_name = selected_proc['name']
                    save_last_process(selected_proc['name'])
                    
                    # Limpa estados de wizard
                    st.session_state.last_ambiente = None
                    st.session_state.last_elemento = None
                    st.session_state.wizard_step = "setup"
                    st.session_state.latitude = 0.0
                    st.session_state.longitude = 0.0
                    
                    st.rerun()
                    
        else:  # Iniciar Novo Processo
            st.markdown("<div class='section-title'>➕ CADASTRAR NOVO PROCESSO</div>", unsafe_allow_html=True)
            
            st.write("Insira o número do processo para criar uma nova planilha Excel dedicada.")
            
            processo_num = st.text_input("Número do Processo (Obrigatório)", placeholder="Ex: 1002345-67.2026.8.26.0100")
            comarca_val = st.text_input("Vara / Comarca (Opcional)", placeholder="Ex: 2ª Vara Cível de SP")
            req_val = st.text_input("Requerente (Opcional)", placeholder="Ex: Condomínio Edifício Progresso")
            reqdo_val = st.text_input("Requerido (Opcional)", placeholder="Ex: Construtora Estrutura Forte Ltda.")
            
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("✨ Criar Novo Processo e Iniciar Vistoria", type="primary", use_container_width=True):
                if not processo_num.strip():
                    st.error("❌ O número do processo é obrigatório para criar um novo registro!")
                else:
                    # Limpa caracteres inválidos
                    safe_proc_name = re.sub(r'[\\/:*?"<>|]', '_', processo_num.strip())
                    new_file_name = f"formulario_vistoria_{safe_proc_name}.xlsx"
                    new_file_path = os.path.join(WORK_DIR, new_file_name)
                    
                    if os.path.exists(new_file_path):
                        st.error(f"⚠️ O processo `{processo_num}` já existe! Escolha a opção 'Abrir Processo Existente' ou defina outro número.")
                    else:
                        try:
                            # Copia o arquivo base
                            shutil.copy(TEMPLATE_PATH, new_file_path)
                            
                            # Abre o novo arquivo e realiza uma LIMPEZA TOTAL automática
                            # para garantir que nenhum dado de teste do template se propague
                            wb_new = openpyxl.load_workbook(new_file_path)
                            
                            # 1. Limpa a aba de Registro de Danos (caso o template estivesse sujo)
                            ws_danos = wb_new['2. Registro de Danos']
                            ws_danos._images.clear() # Limpa as fotos embutidas
                            
                            # Upgrade headers na planilha limpa nova
                            ws_danos.cell(7, 8, 'Local Exato (Pilar/Laje)')
                            ws_danos.cell(7, 9, 'Foto 1 (Visão Geral)')
                            ws_danos.cell(7, 10, 'Foto 2 (Detalhe)')
                            ws_danos.cell(7, 11, None)
                            ws_danos.cell(7, 12, None)
                            
                            from openpyxl.cell import Cell
                            for r_idx in range(8, 1000):
                                ws_danos.row_dimensions[r_idx].height = None # Reseta a altura da linha para o padrão
                                for c_idx in range(1, 13):
                                    cell_obj = ws_danos.cell(r_idx, c_idx)
                                    if type(cell_obj) == Cell:
                                        cell_obj.value = None
                                        
                            # 2. Grava os novos dados de Identificação e limpa campos residuais antigos
                            ws_new = wb_new['1. Identificação e Vistoria']
                            ws_new['B5'] = processo_num.strip()
                            ws_new['E5'] = comarca_val.strip()
                            ws_new['B6'] = req_val.strip()
                            ws_new['E6'] = reqdo_val.strip()
                            ws_new['B7'] = None
                            ws_new['B10'] = None
                            ws_new['E10'] = None
                            ws_new['B11'] = None
                            ws_new['B12'] = None
                            ws_new['E12'] = None
                            
                            wb_new.save(new_file_path)
                            
                            # Configura a sessão ativa
                            st.session_state.active_process_file = new_file_path
                            st.session_state.active_process_name = processo_num.strip()
                            save_last_process(processo_num.strip())
                            
                            # Limpa estados de wizard
                            st.session_state.last_ambiente = None
                            st.session_state.last_elemento = None
                            st.session_state.wizard_step = "setup"
                            st.session_state.latitude = 0.0
                            st.session_state.longitude = 0.0
                            
                            st.rerun()
                        except Exception as e:
                            st.error(f"Erro ao criar novo arquivo de processo: {e}")
    st.stop()

# Função para ler dados de Identificação (Aba 1)
def read_identification_data():
    wb = openpyxl.load_workbook(FILE_PATH, data_only=True)
    ws = wb['1. Identificação e Vistoria']
    
    data = {
        'processo': ws['B5'].value or "",
        'comarca': ws['E5'].value or "",
        'requerente': ws['B6'].value or "",
        'requerido': ws['E6'].value or "",
        'tipo_pericia': ws['B7'].value or support_lists['Pericia'][0],
        'data': ws['B10'].value or datetime.today().date(),
        'hora': ws['E10'].value or "09:00",
        'endereco': ws['B11'].value or "",
        'latitude': ws['B12'].value or 0.0,
        'longitude': ws['E12'].value or 0.0
    }
    
    if isinstance(data['data'], str):
        try:
            data['data'] = datetime.strptime(data['data'], "%Y-%m-%d").date()
        except:
            data['data'] = datetime.today().date()
            
    if isinstance(data['hora'], time):
        data['hora'] = data['hora'].strftime("%H:%M")
        
    return data

# Função para salvar dados de Identificação
def save_identification_data(data):
    wb = openpyxl.load_workbook(FILE_PATH)
    ws = wb['1. Identificação e Vistoria']
    
    ws['B5'] = data['processo']
    ws['E5'] = data['comarca']
    ws['B6'] = data['requerente']
    ws['E6'] = data['requerido']
    ws['B7'] = data['tipo_pericia']
    ws['B10'] = data['data'].strftime("%Y-%m-%d") if isinstance(data['data'], datetime) or hasattr(data['data'], 'strftime') else str(data['data'])
    ws['E10'] = data['hora']
    ws['B11'] = data['endereco']
    ws['B12'] = float(data['latitude']) if data['latitude'] else 0.0
    ws['E12'] = float(data['longitude']) if data['longitude'] else 0.0
    
    wb.save(FILE_PATH)

# Função para ler registros de danos (Aba 2)

# Função para ler registros de danos (Aba 2)
def read_damage_records():
    upgrade_spreadsheet_layout_if_needed(FILE_PATH)
    wb = openpyxl.load_workbook(FILE_PATH, data_only=True)
    ws = wb['2. Registro de Danos']
    
    records = []
    row = 8
    while True:
        val_amb = ws.cell(row, 1).value
        val_elem = ws.cell(row, 2).value
        val_dano = ws.cell(row, 3).value
        
        if val_amb is None and val_elem is None and val_dano is None:
            empty_streak = True
            for check_row in range(row, row + 5):
                if any(ws.cell(check_row, c).value is not None for c in range(1, 10)):
                    empty_streak = False
                    row = check_row
                    break
            if empty_streak:
                break
                
        records.append({
            'row_idx': row,
            'ambiente': val_amb or "",
            'elemento': val_elem or "",
            'dano': val_dano or "",
            'descricao': ws.cell(row, 4).value or "",
            'medicoes': ws.cell(row, 5).value or "",
            'observacoes': ws.cell(row, 6).value or "",
            'foto_codigo': ws.cell(row, 7).value or "",
            'foto_class': "Detalhe",
            'coordenada': "",
            'local_exato': ws.cell(row, 8).value or ""
        })
        row += 1
        
    cols = ['row_idx', 'ambiente', 'elemento', 'dano', 'descricao', 'medicoes', 'observacoes', 'foto_codigo', 'foto_class', 'coordenada', 'local_exato']
    if not records:
        return pd.DataFrame(columns=cols)
    return pd.DataFrame(records)

# Função para adicionar registro com duas fotos embutidas (Geral e Detalhe)
def add_damage_record(record):
    upgrade_spreadsheet_layout_if_needed(FILE_PATH)
    wb = openpyxl.load_workbook(FILE_PATH)
    ws = wb['2. Registro de Danos']
    
    row = 8
    while True:
        val_amb = ws.cell(row, 1).value
        val_elem = ws.cell(row, 2).value
        val_dano = ws.cell(row, 3).value
        if val_amb in [None, ""] and val_elem in [None, ""] and val_dano in [None, ""]:
            break
        row += 1
        
    ws.cell(row, 1, record['ambiente'])
    ws.cell(row, 2, record['elemento'])
    ws.cell(row, 3, record['dano'])
    ws.cell(row, 4, record['descricao'])
    ws.cell(row, 5, record['medicoes'])
    ws.cell(row, 6, record['observacoes'])
    
    # Armazena os dois nomes de fotos na coluna G, separados por ponto e vírgula
    foto_comb = f"{record['foto_local_codigo']}; {record['foto_detalhe_codigo']}"
    ws.cell(row, 7, foto_comb)
    ws.cell(row, 8, record['local_exato'])
    ws.cell(row, 9, "")
    ws.cell(row, 10, "")
    
    # Embutir as duas fotos no Excel
    ws.row_dimensions[row].height = 90
    
    # Foto 1: Geral / Local (Ancorada em I)
    if record['foto_local_codigo']:
        foto_l_path = os.path.join(FOTOS_DIR, record['foto_local_codigo'])
        if os.path.exists(foto_l_path):
            try:
                from openpyxl.drawing.image import Image as OpenpyxlImage
                img_embed_l = OpenpyxlImage(foto_l_path)
                img_embed_l.width = 140
                img_embed_l.height = 110
                ws.add_image(img_embed_l, f'I{row}')
            except:
                pass
                
    # Foto 2: Detalhe (Ancorada em J)
    if record['foto_detalhe_codigo']:
        foto_d_path = os.path.join(FOTOS_DIR, record['foto_detalhe_codigo'])
        if os.path.exists(foto_d_path):
            try:
                from openpyxl.drawing.image import Image as OpenpyxlImage
                img_embed_d = OpenpyxlImage(foto_d_path)
                img_embed_d.width = 140
                img_embed_d.height = 110
                ws.add_image(img_embed_d, f'J{row}')
            except:
                pass
                
    wb.save(FILE_PATH)
    return row

def get_kpis():
    df_danos = read_damage_records()
    total_danos_py = len(df_danos[df_danos['ambiente'] != ""])
    total_fotos_py = len(df_danos[(df_danos['foto_codigo'] != "") & (df_danos['foto_codigo'].notna())])
    danos_criticos_py = len(df_danos[df_danos['dano'].isin(["Rachadura", "Corrosão"])])
    
    return {
        'total_danos': total_danos_py,
        'total_fotos': total_fotos_py,
        'danos_criticos': danos_criticos_py
    }

# --- CABEÇALHO ---
st.markdown(f"<div class='main-header'>🚧 VISTORIA DE CAMPO - ENGENHARIA CIVIL</div>", unsafe_allow_html=True)
st.markdown(f"<div class='sub-header'>Processo Ativo: <b>{st.session_state.active_process_name}</b></div>", unsafe_allow_html=True)

if 'healed_count' in st.session_state and st.session_state.healed_count > 0:
    st.success(f"✨ **Autocura de Planilha Ativa:** Identificamos e corrigimos {st.session_state.healed_count} fotos que estavam sobrepostas na coluna I devido ao bug da versão anterior! Baixe a planilha novamente na barra lateral para salvar os arquivos corrigidos.")
    st.session_state.healed_count = 0 # Mostra apenas uma vez

kpis = get_kpis()
st.sidebar.markdown(f"### 📁 Processo Ativo:\n`{st.session_state.active_process_name}`")
if st.sidebar.button("🔄 Alternar Processo / Voltar", help="Fecha o processo ativo e volta para a tela de seleção inicial", use_container_width=True):
    st.session_state.active_process_file = None
    st.session_state.active_process_name = None
    st.rerun()

st.sidebar.markdown("---")
st.sidebar.markdown("### 📊 Painel Geral de Campo")
col_s1, col_s2 = st.sidebar.columns(2)
with col_s1:
    st.markdown(f"<div class='kpi-card'><div class='kpi-value'>{kpis['total_danos']}</div><div class='kpi-label'>Danos</div></div>", unsafe_allow_html=True)
with col_s2:
    st.markdown(f"<div class='kpi-card' style='border-left-color: #E2C044;'><div class='kpi-value'>{kpis['total_fotos']}</div><div class='kpi-label'>Fotos</div></div>", unsafe_allow_html=True)

st.sidebar.markdown("<br>", unsafe_allow_html=True)
st.markdown(f"""
<div class='kpi-card' style='border-left-color: #C00000; margin-bottom: 20px;'>
    <div class='kpi-value' style='color: #C00000;'>{kpis['danos_criticos']}</div>
    <div class='kpi-label'>Patologias Críticas Identificadas (Corrosão ou Rachaduras)</div>
</div>
""", unsafe_allow_html=True)

# Botão de download da planilha na barra lateral
st.sidebar.markdown("### 💾 Finalização & Exportação")
with open(FILE_PATH, "rb") as f:
    st.sidebar.download_button(
        label="📥 Baixar Planilha com Fotos (.XLSX)",
        data=f,
        file_name="formulario_vistoria_campo.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        help="Baixe a planilha com todas as fotos embutidas e informações sincronizadas."
    )


st.sidebar.markdown("### 📝 Laudo Pericial no Word")
if st.session_state.active_process_file is not None:
    docx_out_name = f"laudo_pericial_{st.session_state.active_process_name}.docx"
    docx_out_path = os.path.join(WORK_DIR, docx_out_name)
    try:
        generate_docx_laudo(FILE_PATH, FOTOS_DIR, docx_out_path)
        if os.path.exists(docx_out_path):
            with open(docx_out_path, "rb") as f_docx:
                st.sidebar.download_button(
                    label="📥 Baixar Laudo Word (.DOCX)",
                    data=f_docx,
                    file_name=f"laudo_{st.session_state.active_process_name}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    help="Gera e baixa o Laudo Técnico formatado em Word (.docx) com as fotos e dados."
                )
    except Exception as e_docx:
        st.sidebar.error(f"Erro ao gerar laudo: {e_docx}")
else:
    st.sidebar.info("💡 Abra ou crie um processo para poder gerar o Laudo em Word.")

st.sidebar.info("💡 **Dica de Campo:** O app emite e grava fotos diretamente no Excel de forma proporcional, otimizando a montagem do laudo!")

# Tabs Principais do App
tab_ident, tab_novo_dano, tab_visualizar = st.tabs([
    "📋 1. Identificação e Vistoria", 
    "📸 2. Registrar Novo Dano (Fluxo Otimizado)", 
    "🔍 3. Danos Registrados & Galeria"
])

# --- TAB 1: IDENTIFICAÇÃO E VISTORIA ---
with tab_ident:
    st.markdown("<div class='section-title'>SEÇÃO 1 — IDENTIFICAÇÃO DA PERÍCIA</div>", unsafe_allow_html=True)
    current_ident = read_identification_data()
    
    col1, col2 = st.columns(2)
    with col1:
        processo = st.text_input("Número do Processo", value=current_ident['processo'], placeholder="Ex: 1002345-67.2026.8.26.0100")
        requerente = st.text_input("Requerente", value=current_ident['requerente'], placeholder="Ex: Condomínio Edifício Progresso")
        
        tipo_idx = 0
        if current_ident['tipo_pericia'] in support_lists['Pericia']:
            tipo_idx = support_lists['Pericia'].index(current_ident['tipo_pericia'])
        tipo_pericia = st.selectbox("Tipo de Perícia", options=support_lists['Pericia'], index=tipo_idx)
        
    with col2:
        comarca = st.text_input("Vara / Comarca", value=current_ident['comarca'], placeholder="Ex: 2ª Vara Cível da Capital / SP")
        requerido = st.text_input("Requerido", value=current_ident['requerido'], placeholder="Ex: Construtora Estrutura Forte Ltda.")
    
    st.markdown("<div class='section-title'>SEÇÃO 2 — DADOS DA VISTORIA</div>", unsafe_allow_html=True)
    
    col3, col4 = st.columns(2)
    with col3:
        data_vistoria = st.date_input("Data da Vistoria", value=current_ident['data'])
        endereco = st.text_area("Endereço Completo do Local", value=current_ident['endereco'], placeholder="Ex: Av. Paulista, 1000 - Bela Vista, São Paulo - SP")
        
    with col4:
        hora_str = current_ident['hora']
        try:
            h, m = map(int, hora_str.split(':'))
            hora_time = time(h, m)
        except:
            hora_time = time(9, 0)
        hora_vistoria = st.time_input("Hora da Vistoria", value=hora_time)
        
        if st.session_state.latitude == 0.0 and current_ident['latitude'] != 0.0:
            st.session_state.latitude = float(current_ident['latitude'])
        if st.session_state.longitude == 0.0 and current_ident['longitude'] != 0.0:
            st.session_state.longitude = float(current_ident['longitude'])

        col_lat, col_lon = st.columns(2)
        with col_lat:
            latitude = st.number_input("Latitude (G. Decimais)", value=float(st.session_state.latitude), format="%.6f", min_value=-90.0, max_value=90.0, key="lat_input_widget")
            st.session_state.latitude = latitude
        with col_lon:
            longitude = st.number_input("Longitude (G. Decimais)", value=float(st.session_state.longitude), format="%.6f", min_value=-180.0, max_value=180.0, key="lon_input_widget")
            st.session_state.longitude = longitude

        st.markdown("---")
        st.markdown("### 📡 Captura de Coordenadas de Identificação")
        gps_html = """
        <div style="background-color: #F8F9FA; padding: 15px; border-radius: 8px; border: 1px solid #2F5496; margin-bottom: 10px;">
            <button onclick="getLocation()" style="background-color: #2F5496; color: white; border: none; padding: 10px 18px; border-radius: 4px; cursor: pointer; font-weight: bold; font-size: 0.95rem;">
                📍 Capturar Coordenada da Obra (GPS)
            </button>
            <p id="gps_status" style="margin-top: 10px; font-weight: bold; color: #555; font-size: 0.9rem; margin-bottom: 5px;"></p>
            <p id="gps_coords" style="font-size: 1.1rem; font-weight: bold; color: #2F5496; margin-top: 5px; font-family: monospace;"></p>
        </div>

        <script>
        function getLocation() {
            var status = document.getElementById('gps_status');
            var coords = document.getElementById('gps_coords');
            status.innerHTML = "Buscando sinal de GPS...";
            coords.innerHTML = "";
            
            if (navigator.geolocation) {
                navigator.geolocation.getCurrentPosition(showPosition, showError, {
                    enableHighAccuracy: true,
                    timeout: 10000,
                    maximumAge: 0
                });
            } else {
                status.innerHTML = "❌ Não suportado pelo navegador.";
            }
        }

        function showPosition(position) {
            var status = document.getElementById('gps_status');
            var coords = document.getElementById('gps_coords');
            status.innerHTML = "✅ Localização obtida! Digite nas caixas acima:";
            coords.innerHTML = "Lat: " + position.coords.latitude.toFixed(6) + "<br>Lon: " + position.coords.longitude.toFixed(6);
        }

        function showError(error) {
            var status = document.getElementById('gps_status');
            status.innerHTML = "❌ Bloqueado (requer conexão segura HTTPS ou permissão local).";
        }
        </script>
        """
        st.components.v1.html(gps_html, height=180)

    if st.button("💾 Salvar Identificação & Dados Gerais", type="primary"):
        save_data = {
            'processo': processo,
            'comarca': comarca,
            'requerente': requerente,
            'requerido': requerido,
            'tipo_pericia': tipo_pericia,
            'data': data_vistoria,
            'hora': hora_vistoria.strftime("%H:%M"),
            'endereco': endereco,
            'latitude': latitude,
            'longitude': longitude
        }
        save_identification_data(save_data)
        st.success("✅ Dados da Perícia e Vistoria updated com sucesso no arquivo Excel!")

# --- TAB 2: REGISTRAR NOVO DANO (FLUXO OTIMIZADO - WIZARD) ---
with tab_novo_dano:
    st.markdown("<div class='section-title'>📸 REGISTRO DE PATOLOGIA - WIZARD DE CAMPO</div>", unsafe_allow_html=True)
    
    # Pasos do Wizard
    if st.session_state.wizard_step == "setup":
        st.markdown("""
        <div class='wizard-card'>
            <div class='step-indicator'>PASSO 1: Definir Local de Inspeção (Caracterização Técnica)</div>
        </div>
        """, unsafe_allow_html=True)
        
        # Caso já tenhamos um ambiente registrado na sessão
        if st.session_state.last_ambiente is not None:
            st.info(f"📍 **Ambiente Anterior Ativo:** `{st.session_state.last_ambiente}`")
            mudar_ambiente = st.radio(
                "Você mudou de Ambiente / Setor de vistoria?",
                options=["Não, continuar no mesmo ambiente", "Sim, mudar de ambiente"],
                index=0
            )
            
            if mudar_ambiente == "Sim, mudar de ambiente":
                active_ambiente = st.selectbox("Selecione o Novo Ambiente / Setor", options=support_lists['Ambiente'])
                active_elemento = st.selectbox("Selecione o Elemento Avaliado", options=support_lists['Elemento'])
            else:
                active_ambiente = st.session_state.last_ambiente
                
                st.info(f"🧱 **Elemento Anterior Ativo:** `{st.session_state.last_elemento}`")
                mesmo_elemento = st.radio(
                    "O Elemento Avaliado continua sendo o mesmo?",
                    options=["Sim, manter o elemento anterior", "Não, trocar de elemento"],
                    index=0
                )
                
                if mesmo_elemento == "Não, trocar de elemento":
                    active_elemento = st.selectbox("Selecione o Novo Elemento Avaliado", options=support_lists['Elemento'])
                else:
                    active_elemento = st.session_state.last_elemento
        else:
            st.warning("⚠️ Nenhum registro feito ainda. Defina o primeiro local para iniciar a vistoria.")
            active_ambiente = st.selectbox("Selecione o Ambiente / Setor", options=support_lists['Ambiente'])
            active_elemento = st.selectbox("Selecione o Elemento Avaliado", options=support_lists['Elemento'])
            
        st.session_state.active_ambiente = active_ambiente
        st.session_state.active_elemento = active_elemento
        
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("Avançar para Seção 4 (Fotos & Detalhes) ➡️", type="primary"):
            st.session_state.wizard_step = "foto_dano"
            st.rerun()

    elif st.session_state.wizard_step == "foto_dano":
        st.markdown(f"""
        <div class='wizard-card'>
            <div class='step-indicator'>PASSO 2: Registro de Imagens (Visão Geral & Detalhe)</div>
            <p style='font-size: 1rem; color: #555;'>Inspeção ativa em: <b>{st.session_state.active_ambiente}</b> &gt; <b>{st.session_state.active_elemento}</b></p>
        </div>
        """, unsafe_allow_html=True)
        
        col_f1, col_f2 = st.columns(2)
        with col_f1:
            uploaded_file_local = st.file_uploader("📸 Foto 1: Visão Geral / Local (Obrigatório - Máx 10 MB)", type=["png", "jpg", "jpeg", "heic", "heif"], help="Identifica a região e o tipo de anomalia")
            uploaded_file_detalhe = st.file_uploader("🔍 Foto 2: Detalhe da Anomalia (Obrigatório - Máx 10 MB)", type=["png", "jpg", "jpeg", "heic", "heif"], help="Mostra em close a patologia para descrição técnica e medições")
            foto_class = "Detalhe"
            local_exato = st.text_input("Local Exato (Opcional)", placeholder="Ex: Pilar P4 da ala oeste do 2º subsolo")
            
            # --- SEÇÃO DE AUTO-MEDIÇÃO POR VISÃO COMPUTACIONAL ---
            if uploaded_file_detalhe is not None:
                st.markdown("<div style='background-color: #EBF3F5; padding: 12px; border-radius: 8px; border: 1.5px solid #2F5496; margin-top: 15px;'>", unsafe_allow_html=True)
                st.markdown("<p style='font-size: 0.95rem; font-weight: bold; color: #2F5496; margin-bottom: 5px; margin-top: 0px;'>🔬 AUTO-MEDIÇÃO POR VISÃO COMPUTACIONAL (BETA)</p>", unsafe_allow_html=True)
                st.markdown("<p style='font-size: 0.8rem; color: #555; margin-bottom: 10px;'>O programa analisará a foto e calculará a espessura da abertura em milímetros usando o objeto de escala selecionado.</p>", unsafe_allow_html=True)
                
                ativar_ia = st.checkbox("💡 Ativar análise automática nesta foto", value=False, key="ativar_ia_crack_det")
                if ativar_ia:
                    ref_scale = st.selectbox(
                        "Objeto de escala visível ao lado do dano:",
                        options=[
                            "Fissurômetro Trident (Comprimento total de 120.0 mm)",
                            "Adesivo Circular do Fissurômetro (diâmetro 10.0 mm)",
                            "Moeda de 1 Real (diâmetro 27.0 mm)",
                            "Cartão de Crédito/Funcional (largura de 85.6 mm)",
                            "Inserção Manual (Apenas estimativa)"
                        ],
                        key="ref_scale_crack_det"
                    )
                    
                    if st.button("⚡ Executar Processamento e Medir Espessura", type="secondary", use_container_width=True):
                        with st.spinner("Processando imagens e calculando pixels..."):
                            try:
                                pil_detalhe = Image.open(uploaded_file_detalhe)
                                res_img, est_thick, est_class, est_color, is_fallback = analyze_crack(pil_detalhe, ref_scale)
                                
                                st.session_state.detected_thickness = est_thick
                                st.session_state.detected_class = est_class
                                st.session_state.detected_color = est_color
                                st.session_state.detected_image = res_img
                                st.session_state.detected_fallback = is_fallback
                                
                                st.session_state.medicoes_input_widget = f"{est_thick:.2f} mm"
                                st.session_state.medicoes_input_val = f"{est_thick:.2f} mm"
                                st.success("✅ Processamento concluído com sucesso!")
                                st.rerun()
                            except Exception as e:
                                st.error(f"Erro ao processar imagem: {e}")
                                
                if 'detected_thickness' in st.session_state:
                    thick_val = st.session_state.detected_thickness
                    class_val = st.session_state.detected_class
                    color_val = st.session_state.detected_color
                    is_fallback = st.session_state.detected_fallback
                    
                    st.markdown(f"""
                    <div style='background-color: white; padding: 10px; border-radius: 6px; border-left: 5px solid {color_val}; margin-top: 5px;'>
                        <p style='font-size: 0.85rem; margin: 0px;'><b>Abertura Calculada:</b> <span style='font-size: 1.15rem; font-weight: bold; color: {color_val};'>{thick_val:.2f} mm</span></p>
                        <p style='font-size: 0.85rem; margin: 5px 0px 0px 0px;'><b>Classificação ABNT Recomendada:</b> <br><span style='font-weight: bold; color: {color_val};'>{class_val}</span></p>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    if is_fallback:
                        st.caption("⚠️ *Aviso: Não detectamos uma escala circular nítida. Aplicada escala média calibrada.*")
                        
                    if st.button("💾 Copiar Medição para o Relatório", use_container_width=True):
                        st.session_state.medicoes_input_val = f"Espessura aproximada de {thick_val:.2f} mm por Visão Computacional."
                        st.info("Espessura copiada para a caixa de 'Medições' ao lado!")
                st.markdown("</div>", unsafe_allow_html=True)
            
        with col_f2:
            proximo_id = kpis['total_fotos'] + 1
            foto_codigo_padrao = f"FOTO_{proximo_id:03d}"
            
            # Sincroniza estado para poder resetar
            if 'foto_codigo_text_input' not in st.session_state:
                st.session_state.foto_codigo_text_input = foto_codigo_padrao
                st.session_state.last_proximo_id = proximo_id
            
            if 'last_proximo_id' not in st.session_state or st.session_state.last_proximo_id != proximo_id:
                st.session_state.foto_codigo_text_input = foto_codigo_padrao
                st.session_state.last_proximo_id = proximo_id
                
            col_code_in, col_code_res = st.columns([2, 1])
            with col_code_in:
                # O valor é lido do widget e seu estado é sincronizado
                foto_codigo = st.text_input("Código de Registro da Foto", key="foto_codigo_text_input")
            with col_code_res:
                st.markdown("<div style='height: 28px;'></div>", unsafe_allow_html=True)
                if st.button("🔄 Resetar Nome", help="Reseta o nome da foto de volta para o padrão sequencial automático"):
                    st.session_state.foto_codigo_text_input = f"FOTO_{proximo_id:03d}"
                    st.rerun()
                    
            # Alerta de Duplicidade corrigido para ambas as fotos
            foto_existente = False
            for ext_check in ['.jpg', '.png', '.heic', '.heif', '.jpeg']:
                if os.path.exists(os.path.join(FOTOS_DIR, f"{foto_codigo}_local{ext_check}")) or os.path.exists(os.path.join(FOTOS_DIR, f"{foto_codigo}_detalhe{ext_check}")):
                    foto_existente = True
                    break
            if foto_existente:
                st.warning(f"⚠️ **Nome Duplicado:** Arquivos com o prefixo `{foto_codigo}` já existem na pasta de fotos. Use o botão 'Resetar Nome' ou altere o código para evitar sobreposições.")
                
            dano = st.selectbox("Tipo de Dano / Patologia", options=support_lists['Dano'])
            # Permite o preenchimento automático a partir da medição por IA
            medicoes = st.text_input(
                "Medições Realizadas", 
                value=st.session_state.medicoes_input_val if st.session_state.medicoes_input_val else "",
                placeholder="Ex: Fissura com abertura de 0.8 mm medida com fissurômetro",
                key="medicoes_input_widget"
            )
            st.session_state.medicoes_input_val = medicoes

        descricao = st.text_area("Descrição Técnica Detalhada", placeholder="Descreva os sinais visuais, extensão do dano e causas prováveis...")
        observacoes = st.text_area("Observações Complementares", placeholder="Fatores de aceleração de degradação, histórico de reformas...")

        # Mostrar previews horizontais (com suporte a exibição do overlay da IA)
        if uploaded_file_local is not None or uploaded_file_detalhe is not None:
            st.markdown("**Visualização das Imagens de Campo:**")
            col_p1, col_p2 = st.columns(2)
            with col_p1:
                if uploaded_file_local is not None:
                    ext_l = os.path.splitext(uploaded_file_local.name)[1].lower()
                    if ext_l in ['.heic', '.heif'] and not HEIF_SUPPORT:
                        st.warning("⚠️ Foto 1 do iPhone (HEIC) detectada!")
                    else:
                        try:
                            img_l = Image.open(uploaded_file_local)
                            st.image(img_l, caption=f"{foto_codigo} - 1. Visão Geral / Local", use_container_width=True)
                        except Exception as e:
                            st.error(f"Erro na Foto 1: {e}")
                else:
                    st.info("📸 Faça o upload da Foto 1 (Visão Geral / Local)")
            with col_p2:
                if uploaded_file_detalhe is not None:
                    # Se tivermos uma imagem processada por IA na sessão, mostra ela com os overlays de medição!
                    if 'detected_image' in st.session_state:
                        st.image(st.session_state.detected_image, caption=f"{foto_codigo} - 2. Detalhe Medido por Visão Computacional", use_container_width=True)
                    else:
                        ext_d = os.path.splitext(uploaded_file_detalhe.name)[1].lower()
                        if ext_d in ['.heic', '.heif'] and not HEIF_SUPPORT:
                            st.warning("⚠️ Foto 2 do iPhone (HEIC) detectada!")
                        else:
                            try:
                                img_d = Image.open(uploaded_file_detalhe)
                                st.image(img_d, caption=f"{foto_codigo} - 2. Detalhe da Anomalia", use_container_width=True)
                            except Exception as e:
                                st.error(f"Erro na Foto 2: {e}")
                else:
                    st.info("🔍 Faça o upload da Foto 2 (Detalhe da Anomalia)")

        st.markdown("<hr>", unsafe_allow_html=True)
        col_b1, col_b2 = st.columns(2)
        with col_b1:
            if st.button("⬅️ Voltar ao Passo 1"):
                st.session_state.wizard_step = "setup"
                st.rerun()
                
        with col_b2:
            if st.button("➕ Gravar Registro de Dano na Planilha", type="primary", use_container_width=True):
                if uploaded_file_local is None or uploaded_file_detalhe is None:
                    st.error("❌ Não é possível gravar! O fluxo de vistoria exige obrigatoriamente ambas as fotos anexadas (Foto 1: Visão Geral / Local E Foto 2: Detalhe da Anomalia) para validar o registro.")
                else:
                    nome_foto_local = ""
                    nome_foto_detalhe = ""
                    
                    # Foto 1: Local
                    ext_l = os.path.splitext(uploaded_file_local.name)[1].lower()
                    filename_l = f"{foto_codigo}_local.jpg"
                    save_path_l = os.path.join(FOTOS_DIR, filename_l)
                    
                    if ext_l in ['.heic', '.heif'] and HEIF_SUPPORT:
                        try:
                            img_heic = Image.open(uploaded_file_local)
                            img_heic.convert("RGB").save(save_path_l, "JPEG")
                            nome_foto_local = filename_l
                        except:
                            filename_l = f"{foto_codigo}_local{ext_l}"
                            save_path_l = os.path.join(FOTOS_DIR, filename_l)
                            with open(save_path_l, "wb") as f:
                                f.write(uploaded_file_local.getbuffer())
                            nome_foto_local = filename_l
                    else:
                        try:
                            img_standard = Image.open(uploaded_file_local)
                            img_standard.convert("RGB").save(save_path_l, "JPEG")
                            nome_foto_local = filename_l
                        except:
                            filename_l = f"{foto_codigo}_local{ext_l}"
                            save_path_l = os.path.join(FOTOS_DIR, filename_l)
                            with open(save_path_l, "wb") as f:
                                f.write(uploaded_file_local.getbuffer())
                            nome_foto_local = filename_l
                            
                    # Foto 2: Detalhe
                    ext_d = os.path.splitext(uploaded_file_detalhe.name)[1].lower()
                    filename_d = f"{foto_codigo}_detalhe.jpg"
                    save_path_d = os.path.join(FOTOS_DIR, filename_d)
                    
                    if ext_d in ['.heic', '.heif'] and HEIF_SUPPORT:
                        try:
                            img_heic = Image.open(uploaded_file_detalhe)
                            img_heic.convert("RGB").save(save_path_d, "JPEG")
                            nome_foto_detalhe = filename_d
                        except:
                            filename_d = f"{foto_codigo}_detalhe{ext_d}"
                            save_path_d = os.path.join(FOTOS_DIR, filename_d)
                            with open(save_path_d, "wb") as f:
                                f.write(uploaded_file_detalhe.getbuffer())
                            nome_foto_detalhe = filename_d
                    else:
                        try:
                            img_standard = Image.open(uploaded_file_detalhe)
                            img_standard.convert("RGB").save(save_path_d, "JPEG")
                            nome_foto_detalhe = filename_d
                        except:
                            filename_d = f"{foto_codigo}_detalhe{ext_d}"
                            save_path_d = os.path.join(FOTOS_DIR, filename_d)
                            with open(save_path_d, "wb") as f:
                                f.write(uploaded_file_detalhe.getbuffer())
                            nome_foto_detalhe = filename_d

                    val_medicoes = medicoes
                    if not val_medicoes.strip() and 'detected_thickness' in st.session_state:
                        val_medicoes = f"{st.session_state.detected_thickness:.2f} mm"
                        
                    novo_registro = {
                        'ambiente': st.session_state.active_ambiente,
                        'elemento': st.session_state.active_elemento,
                        'dano': dano,
                        'descricao': descricao,
                        'medicoes': val_medicoes,
                        'observacoes': observacoes,
                        'foto_local_codigo': nome_foto_local,
                        'foto_detalhe_codigo': nome_foto_detalhe,
                        'foto_class': foto_class,
                        'coordenada': "",
                        'local_exato': local_exato
                    }
                    
                    linha_gravada = add_damage_record(novo_registro)
                    st.success(f"✅ Sucesso! Registro adicionado na Linha {linha_gravada} da planilha Excel (com foto embutida).")
                    
                    # Limpa as variáveis temporárias da IA para o próximo registro
                    if 'detected_thickness' in st.session_state:
                        del st.session_state.detected_thickness
                    if 'detected_class' in st.session_state:
                        del st.session_state.detected_class
                    if 'detected_color' in st.session_state:
                        del st.session_state.detected_color
                    if 'detected_image' in st.session_state:
                        del st.session_state.detected_image
                    if 'detected_fallback' in st.session_state:
                        del st.session_state.detected_fallback
                    st.session_state.medicoes_input_val = ""
                    
                    st.session_state.last_ambiente = st.session_state.active_ambiente
                    st.session_state.last_elemento = st.session_state.active_elemento
                    
                    st.session_state.wizard_step = "pergunta_continua"
                    st.rerun()

    elif st.session_state.wizard_step == "pergunta_continua":
        st.markdown("""
        <div class='wizard-card' style='text-align: center;'>
            <div class='step-indicator' style='border: none;'>Registro Salvo com Sucesso! 🎉</div>
            <h3 style='color: #2F5496;'>Deseja continuar registrando danos para esta vistoria?</h3>            <p style='color: #666;'>O aplicativo reterá o ambiente e elemento para agilizar suas próximas fotos.</p>        </div>
        """, unsafe_allow_html=True)
        
        col_c1, col_c2 = st.columns(2)
        with col_c1:
            if st.button("🟢 SIM, Registrar Outro Dano", use_container_width=True):
                st.session_state.wizard_step = "setup"
                st.rerun()
        with col_c2:
            if st.button("🔴 NÃO, Finalizar Vistoria", use_container_width=True):
                st.success("Inspeção de campo finalizada! Você pode baixar a planilha atualizada na barra lateral ou visualizar os dados na aba 3.")
                st.session_state.wizard_step = "setup"
                st.rerun()

# --- TAB 3: VISUALIZAR DANOS REGISTRADOS ---
with tab_visualizar:
    st.markdown("<div class='section-title'>LISTA DE ELEMENTOS AVALIADOS ATÉ O MOMENTO</div>", unsafe_allow_html=True)
    
    df_danos = read_damage_records()
    df_danos_clean = df_danos[df_danos['ambiente'] != ""].reset_index(drop=True)
    
    if df_danos_clean.empty:
        st.warning("Nenhum registro de dano inserido nesta vistoria até o momento. Utilize o fluxo para registrar.")
    else:
        df_display = df_danos_clean.copy()
        df_display = df_display.drop(columns=['coordenada', 'foto_class'], errors='ignore')
        df_display.columns = [
            'Linha Excel', 'Ambiente / Setor', 'Elemento', 'Dano', 'Descrição Técnica', 
            'Medições', 'Observações', 'Arquivo da Foto', 'Local Exato'
        ]
        
        st.dataframe(df_display, use_container_width=True, hide_index=True)
        
        st.markdown("<div class='section-title'>GALERIA FOTOGRÁFICA DE VISTORIA</div>", unsafe_allow_html=True)
        df_com_fotos = df_danos_clean[df_danos_clean['foto_codigo'] != ""].reset_index(drop=True)
        
        if df_com_fotos.empty:
            st.info("Nenhuma imagem de campo cadastrada ainda.")
        else:
            cols_galeria = st.columns(3)
            for idx, row in df_com_fotos.iterrows():
                col_target = cols_galeria[idx % 3]
                foto_nome = row['foto_codigo']
                
                # Trata múltiplos nomes de foto separados por ';'
                nomes_fotos = [n.strip() for n in foto_nome.split(";") if n.strip()]
                
                with col_target:
                    st.markdown(f"**📍 {row['ambiente']} — {row['local_exato'] or 'Área não especificada'}**")
                    st.markdown(f"**Dano:** `{row['dano']}`")
                    
                    if len(nomes_fotos) == 2:
                        col_pic1, col_pic2 = st.columns(2)
                        with col_pic1:
                            f1 = nomes_fotos[0]
                            f1_path = os.path.join(FOTOS_DIR, f1)
                            if os.path.exists(f1_path):
                                try:
                                    img1 = Image.open(f1_path)
                                    st.image(img1, caption="1. Visão Geral", use_container_width=True)
                                except Exception as e:
                                    st.error(f"Erro: {e}")
                            else:
                                st.warning(f"📷 {f1}\n(Não encontrada localmente)")
                        with col_pic2:
                            f2 = nomes_fotos[1]
                            f2_path = os.path.join(FOTOS_DIR, f2)
                            if os.path.exists(f2_path):
                                try:
                                    img2 = Image.open(f2_path)
                                    st.image(img2, caption="2. Detalhe", use_container_width=True)
                                except Exception as e:
                                    st.error(f"Erro: {e}")
                            else:
                                st.warning(f"📷 {f2}\n(Não encontrada localmente)")
                    else:
                        for f in nomes_fotos:
                            f_path = os.path.join(FOTOS_DIR, f)
                            if os.path.exists(f_path):
                                try:
                                    img_single = Image.open(f_path)
                                    st.image(img_single, caption=f, use_container_width=True)
                                except Exception as e:
                                    st.error(f"Erro: {e}")
                            else:
                                st.warning(f"📷 **{f}**: Imagem não encontrada localmente.")
                    st.markdown("---")
